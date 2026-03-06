#!/usr/bin/env python3
"""
Convert the SPEAR3 LLRF Physical Design Report from Markdown to a professional .docx
with nicely rendered ASCII diagrams as images, formatted tables, and proper styling.
"""

import re
import io
import os
import textwrap
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from PIL import Image, ImageDraw, ImageFont


# ──────────────────────────────────────────────────
# Style and color constants
# ──────────────────────────────────────────────────
TITLE_COLOR = RGBColor(0x00, 0x3C, 0x71)  # Dark navy
HEADING_COLOR = RGBColor(0x1A, 0x4F, 0x8B)  # Medium navy
ACCENT_COLOR = RGBColor(0xCC, 0x55, 0x00)   # Warm accent
TABLE_HEADER_BG = "003C71"  # Dark navy for table headers
TABLE_ALT_ROW = "E8EEF4"    # Light blue for alternating rows
BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"
CODE_FONT = "Consolas"


def setup_styles(doc):
    """Configure document styles for a professional look."""
    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for i in range(1, 5):
        hstyle = doc.styles[f'Heading {i}']
        hstyle.font.name = HEADING_FONT
        hstyle.font.color.rgb = HEADING_COLOR
        hstyle.font.bold = True
        if i == 1:
            hstyle.font.size = Pt(22)
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)
        elif i == 2:
            hstyle.font.size = Pt(17)
            hstyle.paragraph_format.space_before = Pt(18)
            hstyle.paragraph_format.space_after = Pt(8)
        elif i == 3:
            hstyle.font.size = Pt(13)
            hstyle.paragraph_format.space_before = Pt(12)
            hstyle.paragraph_format.space_after = Pt(6)
        elif i == 4:
            hstyle.font.size = Pt(11)
            hstyle.paragraph_format.space_before = Pt(10)
            hstyle.paragraph_format.space_after = Pt(4)


def render_ascii_diagram_to_image(text, title=None):
    """Render an ASCII art diagram into a high-quality PNG image."""
    lines = text.split('\n')
    # Remove leading/trailing empty lines
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    if not lines:
        return None

    # Calculate image size
    max_line_len = max(len(line) for line in lines)
    num_lines = len(lines)

    # Use a monospace font
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", 13)
    except:
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf", 13)
        except:
            font = ImageFont.load_default()

    char_w = 8
    char_h = 17

    # Try to measure actual char size
    try:
        bbox = font.getbbox("M")
        char_w = bbox[2] - bbox[0]
        char_h = bbox[3] - bbox[1] + 4
    except:
        pass

    padding = 30
    title_height = 35 if title else 0
    img_w = max(max_line_len * char_w + 2 * padding, 400)
    img_h = num_lines * char_h + 2 * padding + title_height

    # Create image with light background
    img = Image.new('RGB', (img_w, img_h), '#F7F9FC')
    draw = ImageDraw.Draw(img)

    # Draw a subtle border
    draw.rectangle(
        [4, 4, img_w - 5, img_h - 5],
        outline='#B0C4DE',
        width=2
    )

    # Draw a header bar if title exists
    if title:
        draw.rectangle([4, 4, img_w - 5, title_height + 4], fill='#003C71')
        try:
            title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 13)
        except:
            title_font = font
        draw.text((padding, 12), title, fill='white', font=title_font)

    # Draw the ASCII content
    y = padding + title_height
    for line in lines:
        # Color box-drawing characters in blue for a nice look
        draw.text((padding, y), line, fill='#1A2A3A', font=font)
        y += char_h

    # Save to bytes
    buf = io.BytesIO()
    img.save(buf, format='PNG', dpi=(200, 200))
    buf.seek(0)
    return buf


def add_formatted_run(paragraph, text, bold=False, italic=False, code=False, color=None):
    """Add a formatted run to a paragraph, handling inline markdown."""
    run = paragraph.add_run(text)
    run.font.name = CODE_FONT if code else BODY_FONT
    if code:
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x80, 0x30, 0x00)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_formatting(paragraph, text):
    """Parse inline markdown formatting (bold, italic, code, links) and add runs."""
    # First, handle links: replace [text](url) with just text throughout
    # Then handle bold, italic, code
    
    # Combined pattern for all inline formatting
    # Order: links, inline code, bold+italic, bold, italic
    pattern = r'(\[[^\]]+\]\([^)]+\)|`[^`]+`|\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        # Link: [text](url)
        link_match = re.match(r'^\[([^\]]+)\]\([^)]+\)$', part)
        if link_match:
            add_formatted_run(paragraph, link_match.group(1), color=RGBColor(0x00, 0x56, 0xB3))
        elif part.startswith('`') and part.endswith('`'):
            # Inline code
            code_text = part[1:-1]
            add_formatted_run(paragraph, code_text, code=True)
        elif part.startswith('***') and part.endswith('***'):
            add_formatted_run(paragraph, part[3:-3], bold=True, italic=True)
        elif part.startswith('**') and part.endswith('**'):
            add_formatted_run(paragraph, part[2:-2], bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            add_formatted_run(paragraph, part[1:-1], italic=True)
        else:
            add_formatted_run(paragraph, part)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, header_row, data_rows):
    """Add a professionally formatted table to the document."""
    if not header_row:
        return

    num_cols = len(header_row)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set header row
    hdr = table.rows[0]
    for i, cell_text in enumerate(header_row):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = BODY_FONT
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i >= num_cols:
                break
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            # Parse inline formatting in table cells
            parse_inline_formatting(p, cell_text.strip())
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = BODY_FONT

            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW)

    # Auto-fit columns
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / num_cols)

    doc.add_paragraph()  # spacing after table


def parse_table_block(lines):
    """Parse a markdown table block into header and data rows."""
    if len(lines) < 2:
        return None, None

    def parse_row(line):
        cells = line.strip().strip('|').split('|')
        return [c.strip() for c in cells]

    header = parse_row(lines[0])

    # Skip separator row (the one with ---)
    data_start = 1
    if len(lines) > 1 and re.match(r'^[\s|:-]+$', lines[1]):
        data_start = 2

    data = []
    for line in lines[data_start:]:
        if line.strip() and '|' in line:
            data.append(parse_row(line))

    return header, data


def create_title_page(doc):
    """Create a professional title page."""
    # Add several blank paragraphs for spacing
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('SPEAR3 LLRF Upgrade System')
    run.font.size = Pt(32)
    run.font.color.rgb = TITLE_COLOR
    run.font.name = HEADING_FONT
    run.bold = True

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Physical Design Report')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = HEADING_FONT

    doc.add_paragraph()

    # Horizontal line via a colored paragraph
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run('━' * 50)
    run.font.color.rgb = ACCENT_COLOR
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Meta info
    meta_items = [
        ('Document ID', 'SPEAR3-LLRF-PDR-001'),
        ('Revision', 'R0'),
        ('Date', 'March 3, 2026'),
        ('Author', 'RF Department, SSRL/Accelerator'),
        ('Classification', 'Top-Level System Design Reference'),
    ]

    meta_table = doc.add_table(rows=len(meta_items), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(meta_items):
        row = meta_table.rows[i]
        cell0 = row.cells[0]
        cell1 = row.cells[1]
        cell0.width = Inches(2)
        cell1.width = Inches(3.5)

        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label + ':')
        r0.bold = True
        r0.font.size = Pt(12)
        r0.font.color.rgb = HEADING_COLOR
        r0.font.name = BODY_FONT

        p1 = cell1.paragraphs[0]
        r1 = p1.add_run('  ' + value)
        r1.font.size = Pt(12)
        r1.font.name = BODY_FONT

    # Remove table borders
    for row in meta_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Page break after title
    doc.add_page_break()


def identify_diagram_title(lines, code_start_idx):
    """Try to find a title for a code block from the preceding lines."""
    for i in range(code_start_idx - 1, max(code_start_idx - 5, -1), -1):
        if i < 0:
            break
        line = lines[i].strip()
        if line.startswith('#'):
            return line.lstrip('#').strip()
        if line and not line.startswith('```'):
            return line.rstrip(':')
    return None


def convert_md_to_docx(md_path, output_path):
    """Main conversion function."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    all_lines = content.split('\n')

    doc = Document()
    setup_styles(doc)
    create_title_page(doc)

    # State tracking
    i = 0
    n = len(all_lines)
    in_table = False
    table_lines = []
    skip_meta = True  # Skip the initial metadata block

    # Track which metadata lines to skip (title page already has this info)
    meta_end = 0
    for idx, line in enumerate(all_lines):
        if idx > 20:
            break
        if line.strip() == '---' and idx > 5:
            meta_end = idx + 1
            break

    # Find the second --- to skip the full frontmatter
    dash_count = 0
    for idx, line in enumerate(all_lines):
        if line.strip() == '---':
            dash_count += 1
            if dash_count == 2:
                meta_end = idx + 1
                break
        if idx > 30:
            break

    # Skip the first heading and metadata
    # Find where the actual content begins (after "## Purpose and Scope" or similar)
    i = 0
    # Skip title line
    while i < n and not all_lines[i].strip().startswith('## Purpose'):
        if all_lines[i].strip().startswith('## Table of Contents'):
            break
        if all_lines[i].strip().startswith('## 1.'):
            break
        i += 1
        if i > 50:
            # Fallback: start after second ---
            i = meta_end
            break

    # Actually let's include Purpose and Scope, TOC, etc.
    # Reset to after the header metadata
    i = meta_end if meta_end > 0 else 0
    # Also skip the first # title line since we have title page
    while i < n:
        line = all_lines[i].strip()
        if line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue
        if not line or line == '---':
            i += 1
            continue
        # Skip metadata fields
        if line.startswith('**Document ID**') or line.startswith('**Revision**') or \
           line.startswith('**Date**') or line.startswith('**Author**') or \
           line.startswith('**Classification**'):
            i += 1
            continue
        break

    diagram_counter = 0

    while i < n:
        line = all_lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Code block (diagram)
        if stripped.startswith('```'):
            code_lines = []
            code_lang = stripped[3:].strip()
            i += 1
            while i < n and not all_lines[i].strip().startswith('```'):
                code_lines.append(all_lines[i])
                i += 1
            i += 1  # skip closing ```

            code_text = '\n'.join(code_lines)

            # Determine if this is an ASCII diagram (has box drawing chars) or just code
            has_box_chars = any(c in code_text for c in '┌┐└┘─│├┤┬┴┼▼▲►◄→←↑↓╔╗╚╝═║')
            is_long = len(code_lines) > 5

            if has_box_chars or (is_long and not code_lang):
                # Render as diagram image
                diagram_counter += 1
                title = identify_diagram_title(all_lines, i - len(code_lines) - 2)
                img_buf = render_ascii_diagram_to_image(
                    code_text,
                    title=f"Figure {diagram_counter}: {title}" if title else f"Figure {diagram_counter}"
                )
                if img_buf:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Calculate appropriate width
                    max_line_len = max((len(l) for l in code_lines), default=40)
                    width = min(max_line_len * 0.075, 7.0)  # Scale but cap at 7 inches
                    width = max(width, 4.0)  # Minimum 4 inches
                    run = p.add_run()
                    run.add_picture(img_buf, width=Inches(width))

                    # Add figure caption
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_p.add_run(
                        f"Figure {diagram_counter}" + (f": {title}" if title else "")
                    )
                    cap_run.font.size = Pt(9)
                    cap_run.font.italic = True
                    cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                # Render as code block with grey background
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(cl)
                    run.font.name = CODE_FONT
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Headings
        if stripped.startswith('#'):
            level = 0
            for c in stripped:
                if c == '#':
                    level += 1
                else:
                    break
            heading_text = stripped[level:].strip()
            # Skip the TOC section's link-style entries
            if level <= 4:
                h = doc.add_heading(heading_text, level=min(level, 4))
                # Add a subtle line under H1 and H2
                if level <= 2:
                    line_p = doc.add_paragraph()
                    line_p.paragraph_format.space_before = Pt(0)
                    line_p.paragraph_format.space_after = Pt(4)
                    run = line_p.add_run('─' * 80)
                    run.font.size = Pt(6)
                    run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xDE)
            i += 1
            continue

        # Table detection
        if '|' in stripped and not stripped.startswith('>'):
            table_lines = [stripped]
            i += 1
            while i < n and '|' in all_lines[i].strip() and all_lines[i].strip():
                table_lines.append(all_lines[i].strip())
                i += 1

            header, data = parse_table_block(table_lines)
            if header and data:
                add_table(doc, header, data)
            elif header:
                add_table(doc, header, [])
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            # Add a visual indicator
            run = p.add_run('│ ')
            run.font.color.rgb = ACCENT_COLOR
            run.font.size = Pt(11)
            run.bold = True

            parse_inline_formatting(p, quote_text)
            for run in p.runs[1:]:  # Skip the bar character
                run.font.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            indent_level = len(line) - len(line.lstrip())
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            # Use a bullet character
            bullet_run = p.add_run('• ')
            bullet_run.font.color.rgb = ACCENT_COLOR
            bullet_run.font.size = Pt(11)
            bullet_run.bold = True

            parse_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            num_run = p.add_run(f'{num}. ')
            num_run.font.color.rgb = HEADING_COLOR
            num_run.bold = True

            parse_inline_formatting(p, text)
            i += 1
            continue

        # Indented sub-items (e.g., "  - See: ...")
        if line.startswith('  ') and stripped.startswith('- '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            bullet_run = p.add_run('◦ ')
            bullet_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            bullet_run.font.size = Pt(10)

            parse_inline_formatting(p, text)
            for run in p.runs[1:]:
                run.font.size = Pt(10)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph (possibly a continuation or standalone line)
        # Check for TOC link entries like "1. [Executive Summary](#...)"
        toc_match = re.match(r'^(\d+)\.\s+\[([^\]]+)\]\(#[^)]*\)\s*$', stripped)
        if toc_match:
            # TOC entry - format nicely (no duplication)
            num = toc_match.group(1)
            title = toc_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f'{num}.  {title}')
            run.font.size = Pt(11)
            run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)
        i += 1

    doc.save(output_path)
    print(f"✅ Document saved to: {output_path}")
    return output_path


if __name__ == '__main__':
    md_path = '/tmp/fayaw/spearlegacyLLRF/Designs/0_PHYSICAL_DESIGN_REPORT.md'
    output_path = '/tmp/fayaw/spearlegacyLLRF/Designs/0_PHYSICAL_DESIGN_REPORT.docx'
    convert_md_to_docx(md_path, output_path)
